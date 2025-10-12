from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Caminho da imagem (substitua pelo caminho da sua imagem)
imagem_path = "caminho/para/sua/imagem.jpg"

# Criação do documento conforme ABNT NBR 14724:2024
doc = Document()

# Configurações de margens
for section in doc.sections:
    section.top_margin = Inches(1.18)  # 3 cm
    section.bottom_margin = Inches(0.79)  # 2 cm
    section.left_margin = Inches(1.18)  # 3 cm
    section.right_margin = Inches(0.79)  # 2 cm

# Fonte padrão
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ---------------- CAPA ----------------
doc.add_paragraph("UNIVERSIDADE EXEMPLO", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("NOME FICTÍCIO EXEMPLO", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("\n\nTRABALHO DE AULA – TÍTULO DO TRABALHO EXEMPLO\n\n", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("CIDADE EXEMPLO - UF\n2025", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Quebra de página
doc.add_page_break()

# ---------------- FOLHA DE ROSTO ----------------
doc.add_paragraph("NOME FICTÍCIO EXEMPLO", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("\n\nTRABALHO DE AULA – TÍTULO DO TRABALHO EXEMPLO\n\n", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    "Trabalho apresentado à disciplina de [Nome da Disciplina] do curso de [Nome do Curso] da [Nome da Universidade], "
    "como requisito parcial para avaliação acadêmica.", style='Normal'
).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_paragraph("\n\nCIDADE EXEMPLO - UF\n2025", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Quebra de página
doc.add_page_break()

# ---------------- RESUMO ----------------
titulo_resumo = doc.add_paragraph("RESUMO")
titulo_resumo.alignment = WD_ALIGN_PARAGRAPH.CENTER
titulo_resumo.runs[0].bold = True

resumo = (
    "O presente trabalho tem como objetivo abordar [tema do trabalho], "
    "analisando [aspectos principais do tema]. Além disso, apresenta uma "
    "reflexão sobre [elementos importantes relacionados ao tema]. "
    "A metodologia utilizada baseou-se em [descrição da metodologia]. "
    "Os resultados obtidos demonstram [principais conclusões]."
)
doc.add_paragraph(resumo)
doc.add_paragraph("Palavras-chave: palavra1; palavra2; palavra3; palavra4.")

# Quebra de página
doc.add_page_break()

# ---------------- SUMÁRIO ----------------
doc.add_paragraph("SUMÁRIO", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("1 INTRODUÇÃO .......................................................................................................................... 4")
doc.add_paragraph("2 DESENVOLVIMENTO .................................................................................................................. 5")
doc.add_paragraph("3 CONCLUSÃO ............................................................................................................................. 9")
doc.add_paragraph("REFERÊNCIAS ........................................................................................................................... 10")

doc.add_page_break()

# ---------------- INTRODUÇÃO ----------------
titulo_intro = doc.add_paragraph("1 INTRODUÇÃO")
titulo_intro.runs[0].bold = True

introducao = (
    "Este trabalho aborda [tema central do trabalho], que é considerado um tópico de grande relevância "
    "na área de [área de estudo]. O objetivo principal é [objetivo principal do trabalho]. "
    "A importância deste estudo reside no fato de que [justificativa da importância]. "
    "Para atingir os objetivos propostos, foi utilizada uma metodologia baseada em [metodologia utilizada]. "
    "Este documento está estruturado em seções que abordam [estrutura do trabalho]."
)
doc.add_paragraph(introducao)

# ---------------- DESENVOLVIMENTO ----------------
titulo_desen = doc.add_paragraph("2 DESENVOLVIMENTO")
titulo_desen.runs[0].bold = True

# Subtópicos
doc.add_paragraph("2.1 Primeiro tópico do desenvolvimento", style='List Number')
texto1 = (
    "Aqui você deve desenvolver o primeiro aspecto do seu tema. Apresente conceitos fundamentais, "
    "definições importantes e contextualize o leitor sobre os elementos básicos necessários para "
    "a compreensão do trabalho."
)
doc.add_paragraph(texto1)

doc.add_paragraph("2.2 Segundo tópico do desenvolvimento", style='List Number')
texto2 = (
    "Nesta seção, aprofunde-se em aspectos específicos do tema. Apresente dados, estatísticas, "
    "comparações ou análises que suportem sua argumentação. Use referências bibliográficas "
    "para embasar suas afirmações."
)
doc.add_paragraph(texto2)

doc.add_paragraph("2.3 Terceiro tópico do desenvolvimento", style='List Number')
texto3 = (
    "Desenvolva aqui outro aspecto relevante do tema. Pode incluir estudos de caso, exemplos práticos, "
    "ou análises críticas que contribuam para o entendimento completo do assunto abordado."
)
doc.add_paragraph(texto3)

doc.add_paragraph("2.4 Quarto tópico do desenvolvimento", style='List Number')
texto4 = (
    "Finalize o desenvolvimento abordando implicações, consequências ou aplicações práticas "
    "do tema estudado. Faça conexões entre os diferentes aspectos apresentados nas seções anteriores."
)
doc.add_paragraph(texto4)

# Adiciona imagem (opcional - remova se não precisar)
doc.add_paragraph("Figura 1 – Título explicativo da imagem exemplo", style='Normal')
# Descomente a linha abaixo se quiser incluir uma imagem
# doc.add_picture(imagem_path, width=Inches(5.5))
legenda = doc.add_paragraph("Fonte: Autor exemplo (2025).", style='Normal')
legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("2.5 Considerações finais do desenvolvimento", style='List Number')
texto5 = (
    "Para finalizar o desenvolvimento, é importante sintetizar os principais pontos abordados "
    "e estabelecer conexões entre eles. Esta seção deve preparar o leitor para as conclusões "
    "que serão apresentadas na próxima seção do trabalho."
)
doc.add_paragraph(texto5)

# ---------------- CONCLUSÃO ----------------
titulo_conc = doc.add_paragraph("3 CONCLUSÃO")
titulo_conc.runs[0].bold = True

conclusao = (
    "Com base no desenvolvimento apresentado, pode-se concluir que [principais conclusões do trabalho]. "
    "Os objetivos propostos foram atingidos através de [metodologia utilizada], demonstrando que "
    "[resultados obtidos]. Este estudo contribui para [contribuição do trabalho] e sugere que "
    "[sugestões para trabalhos futuros ou aplicações práticas]."
)
doc.add_paragraph(conclusao)

# ---------------- REFERÊNCIAS ----------------
doc.add_paragraph("REFERÊNCIAS", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
referencia = (
    "AUTOR EXEMPLO, Nome. Título do artigo exemplo. "
    "Nome da Revista Exemplo, v. 1, n. 1, p. 1-10, 2025. DOI: 10.1000/exemplo."
)
doc.add_paragraph(referencia)

# Adicione mais referências conforme necessário
# Exemplo de livro:
referencia2 = (
    "SOBRENOME, Nome do Autor. Título do livro exemplo. Cidade: Editora Exemplo, 2025."
)
doc.add_paragraph(referencia2)

# Salvar arquivo
caminho_docx = "Trabalho_ABNT_Exemplo.docx"
doc.save(caminho_docx)

print(f"Documento salvo como: {caminho_docx}")
print("Lembre-se de substituir todas as informações entre [] pelas suas informações específicas!")
